############################
#UTF-8
#Shawn.Peng@quantatw.com
#Can help you to double confirm info between release_note and BCU
############################
# There are its base functions:
# 1.Get base information in BCU.txt witch show in BIOS R-N(exactly same name).
# 2.Try to get Check Sum from .bin.
# 3.Try to get Sprint(need local build).
# 4.Try to Get Agesa PI / SMU / PSP /… from amdz(if AMD).
# 5.Try to Get ME/RC/GbE… (if Intel).
# 6.Try to get Exteranl Link version.
# 7.Try to get SHA256.
# 8.Try to get FUR.exe version.
# (If 4~6 get fail, it still continue.)
# 9.And finally compare above 1~8 information then save the result.xlsm.
# Note : The BCU.txt / BIOS Release_Note.xlsm must exist and do not opening. The result.xlsm can not opening either(Maybe you run once).
#
# Share point
# xxx\CMIT_BIOS\Tools\compare_BCU_RN\compare_BCU_RN_V{number}.7z
# GitHub 
# https://github.com/yishawnpeng/HpBIOS_ReleaseNote_Check.git
#
############################
from pandas import *    #excel
import sys              #exit don't use os.exit
import shutil           #for copy file (os.rename will remove file)
import docx
from win32com.client import * # GetFileVersion from exe
from lib import *
from tabula import read_pdf         #read pdf form
import openpyxl         #fill coler
from datetime import datetime   # Protect excel
import hashlib                  # Protect excel
import xml.etree.ElementTree as ET # Read System Scope
from collections import Counter    # Read System Scope

version = "10.2"
#support G3/G4 (release note docx)
arg=argparse_function(version)

AMDPlatformDict = {"R24","R26","S25","S27","S29","T25","T26","T27"}
AMDG4PlatformDict = {"Q26","Q27"}
isAMDPlatform = None
isAMDG4Platform = None
AMIPlatformDict = {"U24"}
isAMIPlatform = None
isG4Platform = None
###print("ME is compare form BCU not .bin")

#=========Let user input platform and version
print("v"+version)
goal_platform = input("Input Platform : ")
goal_version = input("Input Version : ")
release_dir = ".\\"+str(goal_platform)+"_"+str(goal_version)
fatherDir = os.getcwd() # father / outside
allDir = os.listdir( fatherDir ) #list
#====Check G4
if "Q" in goal_platform or "P" in goal_platform :
    if goal_platform in AMDG4PlatformDict :
        isAMDG4Platform = True
    else :
        isG4Platform = True
#====Create folder
if os.path.isdir(release_dir):
    os.chdir(release_dir)
    release_dir = os.getcwd()
    os.chdir(fatherDir)
    new_dir = ".\\"+str(goal_platform)+"_"+str(goal_version)+"_checked"
    if not os.path.exists(new_dir) :
        print("Create folder : " + new_dir)
        os.makedirs(new_dir)
        os.chdir(new_dir)
        new_dir = os.getcwd()
        os.chdir(fatherDir)
    else :
        print("Folder already exist : " + new_dir)
        os.chdir(new_dir)
        new_dir = os.getcwd()
        os.chdir(fatherDir)
elif (list( filter( re.compile(fr'(?=.*{re.escape(goal_platform)})(?=.*{re.escape(goal_version)})').match, allDir ) )) : 
    # as AsteroidsR_PV_U23_926100 need to find more
    release_dir = re.compile(fr'(?=.*{re.escape(goal_platform)})(?=.*{re.escape(goal_version)})')
    release_dir = list( filter( release_dir.match, allDir ) )
    for i in release_dir :
        if os.path.isdir(i) and \
            ( "Fv" not in i and "Buildjob" not in i ) :
            release_dir = ".\\"+i
            #print(i+" is dir")
            break
    #print(release_dir)
    os.chdir(release_dir)
    release_dir = os.getcwd()
    os.chdir(fatherDir)
    new_dir = ".\\"+str(goal_platform)+"_"+str(goal_version)+"_checked"
    if not os.path.exists(new_dir) :
        print("Create folder : " + new_dir)
        os.makedirs(new_dir)
        os.chdir(new_dir)
        new_dir = os.getcwd()
        os.chdir(fatherDir)
    else :
        print("Folder already exist : " + new_dir)
        os.chdir(new_dir)
        new_dir = os.getcwd()
        os.chdir(fatherDir)
else :
    print("Can not find : " + release_dir)
    os.system("pause")
    sys.exit()
#====Move file (from outside)
#BCU
bcu_name = re.compile(".*BCU\.txt|.*bcu\.txt") # {any}BCU.txt or {any}bcu.txt
bcu_name = list( filter( bcu_name.match, allDir ) )
if not bcu_name :   #empty
    print("You don't have BCU file or being opened!\nFormat should be \{any\}BCU.txt or \{any\}bcu.txt !")
    os.system("pause")
    sys.exit()
else :
    print("Choose BCU : " + bcu_name[0])
    if not os.path.isfile(new_dir+"\\"+bcu_name[0]) :
        os.rename(fatherDir+"\\"+bcu_name[0],new_dir+"\\"+bcu_name[0])
#AMDZ
amdz_name = re.compile("amdz.*\.txt")
amdz_name = list( filter( amdz_name.match, allDir ) )
if not amdz_name : # empty
    if isAMDPlatform :
        print("You don't have amdz file !\nFormat : amdz\{any\}.txt")
else :
    print("Choose amdz : " + amdz_name[0])
    if not os.path.isfile(new_dir+"\\"+amdz_name[0]) :
        os.rename(fatherDir+"\\"+amdz_name[0],new_dir+"\\"+amdz_name[0])
#External Link
el_name = re.compile(".*External_Link\.txt|.*EL\.txt") # {any}External_Link.txt or {any}EL.txt
el_name = list( filter( el_name.match, allDir ) )
if not el_name :
    print("Can not find \{any\}External_Link.txt!\nFormat : \{any\}External_Link.txt or \{any\}EL.txt")
else :
    print("Choose External Link : " + el_name[0])
    if not os.path.isfile(new_dir+"\\"+el_name[0]) :
        os.rename(fatherDir+"\\"+el_name[0],new_dir+"\\"+el_name[0])
#PDF -- for intel WW info
pdf_name = re.compile(".*_WW\d*_.*\.pdf") # {any}_WW{number}_{any}.pdf
pdf_name = list( filter( pdf_name.match, allDir ) )
if not pdf_name :
    if not isAMDPlatform and not isAMIPlatform :
        print("Can not find \{any\}_WW\{number\}_\{any\}.pdf!\n")
else :
    print("Choose PDF : " + pdf_name[0])
    if not os.path.isfile(new_dir+"\\"+pdf_name[0]) :
        os.rename(fatherDir+"\\"+pdf_name[0],new_dir+"\\"+pdf_name[0])
#ME -- for intel
me_name = re.compile(".*ME.*\.txt") # {any}ME{any}.txt
me_name = list( filter( me_name.match, allDir ) )
if not me_name :
    if not isAMDPlatform and not isAMIPlatform :
        print("Can not find \{any\}ME\{any\}.txt!\n")
else :
    print("Choose ME : " + me_name[0])
    if not os.path.isfile(new_dir+"\\"+me_name[0]) :
        os.rename(fatherDir+"\\"+me_name[0],new_dir+"\\"+me_name[0])
#drivers -- for intel
dri_name = re.compile(".*drivers.*\.txt") # {any}drivers{any}.txt
dri_name = list( filter( dri_name.match, allDir ) )
if not dri_name :
    if not isAMDPlatform and not isAMIPlatform :
        print("Can not find \{any\}drivers\{any\}.txt!\n")
else :
    print("Choose drivers : " + dri_name[0])
    if not os.path.isfile(new_dir+"\\"+dri_name[0]) :
        os.rename(fatherDir+"\\"+dri_name[0],new_dir+"\\"+dri_name[0])
#SystemScope -- for intel
SystemScope_name = re.compile(".*SystemScope.*\.xml") # {any}SystemScope{any}.xml
SystemScope_name = list( filter( SystemScope_name.match, allDir ) )
if not SystemScope_name :
    if not isAMDPlatform and not isAMIPlatform :
        print("Can not find \{any\}SystemScope\{any\}.xml!\n")
else :
    SystemScope_name = SystemScope_name[0]
    print("Choose SystemScope : " + SystemScope_name)
    if not os.path.isfile(new_dir+"\\"+SystemScope_name) :
        os.rename(fatherDir+"\\"+SystemScope_name,new_dir+"\\"+SystemScope_name)
#====Copy file (from release file)
os.chdir(release_dir)
release_all_dir = os.listdir( os.getcwd() )
#Release Note
excelName = re.compile("\w.*Release_Note_\d*\.xlsm|\w.*Release_Note.xlsm|\w.*Release Note_\d*\.xlsm")
if isAMDG4Platform or isG4Platform :
    excelName = re.compile("\w.*release note.docx|\w.*_Release_Notes.docx")
excelName = list( filter( excelName.match, release_all_dir ) )
#print("excelName",excelName)
#print("release_all_dir",release_all_dir)
#print(os.getcwd())
if not excelName :
    if isAMDG4Platform or isG4Platform :
        print("Can not find Release Note!\nFormat :\{any\}release note.docx or \{any\}_Release_Notes.docx")
    else :
        print("Can not find Release Note!\nFormat :\{any\}Release_Note_\{number\}.xlsm or \{any\}Release_Note.xlsm")
    os.system("pause")
    sys.exit()
elif len(excelName) == 1 :
    chooseNote = excelName[0]
else : # multy Note
    chooseStr = "\n".join([f"[{index}] {item}" for index, item in enumerate(excelName)])
    chooseNote = input(f"You have multy ReleaseNote!\nChoose one:\n{chooseStr}\n: ")
    chooseNote = excelName[int(chooseNote)]
    #print("Choose Release Note : " + chooseNote)
shutil.copy(release_dir+"\\"+chooseNote, new_dir+"\\"+chooseNote)
#SHA256
SHA256_file = re.compile(".*\d+_SHA256.txt")
SHA256_file = list( filter( SHA256_file.match, release_all_dir ) )
SHA256isGetOutside=False
if not SHA256_file :
    print("Can not find \{any\}[number]_SHA256.txt !")
    try :
        print("Try outside folder ~ ")
        #os.chdir(fatherDir)
        #print(allDir)
        SHA256_file = re.compile(".*\d+_SHA256.txt")
        SHA256_file = list( filter( SHA256_file.match, allDir ) ) 
        if not SHA256_file :
            print("Can not find \{any\}[number]_SHA256.txt outside also!")
        else :
            print("Choose outside SHA256 : " + SHA256_file[0])
            SHA256isGetOutside = True
            os.rename(fatherDir+"\\"+SHA256_file[0],new_dir+"\\"+SHA256_file[0])
        #os.chdir(release_dir)
    except Exception as e :
        #print(e)
        print("Get outside folder Error : SHA256_file")
        #os.chdir(release_dir)
else :
    print("Choose SHA256 : " + SHA256_file[0])
    shutil.copy(release_dir+"\\"+SHA256_file[0], new_dir+"\\"+SHA256_file[0])
#FUR (from HPFWUPDREC)
information_parser = Dispatch("Scripting.FileSystemObject")
furP = ".\\HPFWUPDREC\\HpFirmwareUpdRec64.exe"
if os.path.isfile(furP) :
    shutil.copy(release_dir+"\\"+furP, new_dir+"\\"+"HpFirmwareUpdRec64.exe")
    furP = "HpFirmwareUpdRec64.exe"
    #furV = information_parser.GetFileVersion(r".\\HPFWUPDREC\\HpFirmwareUpdRec64.exe")
else :
    ##AMI start
    furP=""
    ami_furP = re.compile("U24_\d+.exe")
    ami_furP = list( filter( ami_furP.match, release_all_dir ) )
    if os.path.isfile(ami_furP[0]) :
        shutil.copy(release_dir+"\\"+ami_furP[0], new_dir+"\\"+ami_furP[0])
        ami_furP = ami_furP[0]
        #furV = information_parser.GetFileVersion(platform + "_" + version + ".exe")
    else :
        print("Can not find .\\U24_[number].exe !")
    ##AMI end
    print("Can not find .\\HPFWUPDREC\\HpFirmwareUpdRec64.exe !")
#bin (for checksum from Global/BIOS)
if os.path.isfile(".\\Global\\BIOS\\"+goal_platform+"_"+goal_version+".bin") :
    #AMD
    shutil.copy(release_dir+"\\Global\\BIOS\\"+goal_platform+"_"+goal_version+".bin"\
              , new_dir+"\\"+goal_platform+"_"+goal_version+".bin")
    binFile = goal_platform+"_"+goal_version+".bin"
elif os.path.isfile(".\\Global\\BIOS\\"+goal_platform+"_"+goal_version+"_16.bin") :
    #AMI
    shutil.copy(release_dir+"\\Global\\BIOS\\"+goal_platform+"_"+goal_version+"_16.bin"\
              , new_dir+"\\"+goal_platform+"_"+goal_version+"_16.bin")
    binFile = goal_platform+"_"+goal_version+"_16.bin"
elif os.path.isfile(".\\Global\\BIOS\\"+goal_platform+"_"+goal_version+"_32.bin") :
    #Intel
    shutil.copy(release_dir+"\\Global\\BIOS\\"+goal_platform+"_"+goal_version+"_32.bin"\
              , new_dir+"\\"+goal_platform+"_"+goal_version+"_32.bin")
    binFile = goal_platform+"_"+goal_version+"_32.bin"
else :
    print("Can not find \{platform\}_\{version\}_[|16|32].bin !")
#=========Compare info
#Go to new folder
os.chdir(new_dir)
#Get BCU info
bcu_content=[]
with open(bcu_name[0]) as f:
    for line in f.readlines():
        bcu_content.append(line)
#Get Release Note info
if isAMDG4Platform or isG4Platform :
    rName = chooseNote
    if isAMDG4Platform :
        platform = rName.split("_")[1]
    else :
        platform = rName.split("_")[2]
    if isAMDG4Platform :
        version = rName.split("_")[2]
        version = version.replace(".","")
    else :
        version = rName.split("_")[3]
        version = version.split(" ")[0]
    if goal_platform != str(platform) or goal_version != str(version) :
        print("\nYour INPUT plateform_version is different with geted release note plateform_version!\nYou might ckeck!\n")
    try :
        """ G4
        file  = docx.Document("Scotty_Q26_02.22.00_0001_Release_Notes.docx")
        #intel Sax_PV_Q11_022201 release note.docx
        #print("len",len(file.paragraphs))
        tables=file.tables
        table = tables[1]
        for i in range(0,len(table.rows)) :
            result = table.cell(i,0).text
            r2 = table.cell(i,1).text
            print(result)
            print(r2)
        """
        rRowInfoName = docx.Document(rName)
        table=rRowInfoName.tables[1]
        rRowInfoName=[]
        rRowData=[]
        for i in range(0,len(table.rows)) :
            rRowInfoName.append(table.cell(i,0).text)
            rRowData.append(table.cell(i,1).text)

        #Find Item Range
        startIndex = rRowInfoName.index("System BIOS")
        endIndex = rRowInfoName.index("CHID") #CHID
    except Exception :
        #print(Exception)
        print("Get release note info! May be ceil(sheet) name error.")
        os.system("pause")
        sys.exit()
else :
    rName = chooseNote
    platform = rName.split("_")[2]
    version = rName.split("_")[-1].split(".")[0]
    isR = True if rName.split("_")[0][-1] == "R" else False
    isAMIPlatform = True if goal_platform in AMIPlatformDict else False
    if isAMIPlatform :              # \w.*Release_Note.xlsm
        version = "".join(bcu_content[bcu_content.index("BIOS Revision\n")+1].strip()[1:].split("."))
    isAMDPlatform = True if goal_platform in AMDPlatformDict else False
    if goal_platform != str(platform) or goal_version != str(version) :
        print("\nYour INPUT plateform_version is different with geted release note plateform_version!\nYou might ckeck!\n")
    ##Get item name and info of this time
    try :
        #Item name              :   usecols=[0]
        #Get from Release Note  :   usecols=[1]
        if isAMDPlatform :
            rRowInfoName = read_excel( rName, sheet_name = "AMDPlatformHistory", usecols=[0] )
            rRowData = read_excel( rName, sheet_name = "AMDPlatformHistory", usecols=[1] )
        elif platform in {"U21","U23"} :
            if isR :
                rRowInfoName = read_excel( rName, sheet_name = "IntelPlatformHistory_FY23", usecols=[0] )
                rRowData = read_excel( rName, sheet_name = "IntelPlatformHistory_FY23", usecols=[1] )
            else :
                rRowInfoName = read_excel( rName, sheet_name = "IntelPlatformHistory_FY22", usecols=[0] )
                rRowData = read_excel( rName, sheet_name = "IntelPlatformHistory_FY22", usecols=[1] )
        elif goal_platform in {"U21","U23","U11"} and not platform in {"U21","U23","U11"} :
            # new intel U21/U23/U11
            rRowInfoName = read_excel( rName, sheet_name = "PlatformHistory", usecols=[0] )
            rRowData = read_excel( rName, sheet_name = "PlatformHistory", usecols=[1] )
        # include Intel AMI
        else :
            rRowInfoName = read_excel( rName, sheet_name = "IntelPlatformHistory", usecols=[0] )
            rRowData = read_excel( rName, sheet_name = "IntelPlatformHistory", usecols=[1] )
        rRowInfoName = rRowInfoName[rRowInfoName.columns[0]].tolist()
        #Find Item Range
        startIndex = rRowInfoName.index("System BIOS Version")
        endIndex = rRowInfoName.index("Sprint Release Note") # Sprint Release Note
    except Exception as e :
        print(e)
        print("Get release note info! May be ceil(sheet) name error.")
        os.system("pause")
        sys.exit()
#Get AMDZ info
if amdz_name :
    amdz_content=[]
    with open(amdz_name[0]) as f:
        for line in f.readlines():
            amdz_content.append(line)
#Get SHA256 info
SHA256_content=[]
SHA256UTF16le = True
if SHA256_file :
    try : 
        with open(SHA256_file[0], encoding = "utf-16le") as f:
            for line in f.readlines():
                SHA256_content.append(line)
    except :  # SHA256 is from Outside
        print("SHA256 is not utf16le, try get default.")
        SHA256UTF16le = False
        with open(SHA256_file[0]) as f:
            for line in f.readlines():
                SHA256_content.append(line)
#Get External Link
if el_name:
    el_content = []
    with open(el_name[0]) as f:
        for line in f.readlines():
            el_content.append(line)
#Get PDF
if pdf_name:
    pdfTables = read_pdf(pdf_name[0], pages='all', multiple_tables=True,encoding="ISO-8859-1")
    if not pdfTables[-1].empty :
        pdfTables.pop()
#Get ME
if me_name:
    me_content = []
    with open(me_name[0]) as f:
        for line in f.readlines():
            me_content.append(line)
#Get drivers
if dri_name:
    dri_content = []
    with open(dri_name[0], encoding = "utf-16le") as f:
        for line in f.readlines():
            dri_content.append(line)
#Get System Scope
ssgoals=["Reference Code - MRC","ISHC FW Version","TXT ACM version",\
         "Microcode Version","PMC FW Version","OEM Chipset Init Version"]
gCount = Counter(ssgoals)
if SystemScope_name :
    for event, elem in ET.iterparse(SystemScope_name , events=("start",)):
        if elem.tag == "Item" :
            # try to get "version" or maybe "value"
            if elem.get("Name") in ssgoals :
                #print(elem.get("Name"), " in Goal.")
                tempName = elem.get("Name")
                if gCount[tempName] == 1 :
                    if elem.get("Version") :
                        tempV = elem.get("Version")
                    elif elem.get("Value") :
                        tempV = elem.get("Value")
                    else :
                        print(tempName+" : Can not get from Scope either Value nor Version!")
                        continue
                    #print(tempName, tempV)
                    gCount[tempName] = tempV
                continue
#Create resault.xml
try:
    writer = ExcelWriter(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
    outputFile = []
    outputFile_PlatformHistory = DataFrame( index = rRowInfoName[startIndex:endIndex], \
                                                columns = ["Release Note Info", "Reference Info", "Result"] )
    outputFile.append(outputFile_PlatformHistory) #Sheet No.1
    if isG4Platform or isAMDG4Platform :
        outputFile[0].iloc[:, 0] = rRowData[startIndex:endIndex]
    else :
        outputFile[0].iloc[:, 0] = rRowData[rRowData.columns[0]].tolist()[startIndex:endIndex] 
except Exception as e:
    print(e)
    print("Creat excel fail or \{platform\}result_RN.xlsx being opened!")
    os.system("pause")
    sys.exit()
#common
rRowInfoName = rRowInfoName[startIndex:endIndex]
isAMDBlock = False # for Intel new RN - to check once Microcode
for i in rRowInfoName:
    if type(i) != str : # skip nan
        continue
    try :
        name = i +"\n"
        index = bcu_content.index(name)+1
        if i == "System BIOS Version" :
            temp = bcu_content[index].split()[2]
        elif i =="BIOS Build Version" :
            temp = bcu_content[index].strip()
        else :
            temp = bcu_content[index]
        outputFile[0].at[i, "Reference Info"] = temp
    except Exception :
        if i == "Build Date" and not isAMIPlatform :
            try:
                bdate = bcu_content[bcu_content.index("System BIOS Version\n")+1].split()[-1]
                outputFile[0].at[i, "Reference Info"] = bdate
                continue
            except :
                pass
        elif i == "CHECKSUM" and not isAMIPlatform :
            if binFile :
                with open(binFile, 'rb') as f:
                    content = f.read()
                    binary_sum = sum(bytearray(content))
                    binary_sum = hex(binary_sum & 0xFFFFFFFF)
                    f.close()
                #x need lower, other need upper
                binary_sum = "0x"+binary_sum.split("x")[-1].upper()
                outputFile[0].at[i, "Reference Info"] = binary_sum
            continue
        #elif i == "Sprint":
            #print("Sprint info in local build BCU!")
        elif ( i=="EC/SIO F/W" or i=="SIO FW" or i=="EC/SIO FW" ) and not isAMIPlatform :
            try:
                sio = bcu_content[bcu_content.index("Super I/O Firmware Version\n")+1].split()[-1]
                outputFile[0].at[i, "Reference Info"] = sio
                continue
            except :
                pass
        elif isTypecPD(i) :
            try :
                firstPD = bcu_content[bcu_content.index("USB Type-C Controller(s) Firmware Version:\n")+1].split()[-1]
                secondPD = bcu_content[bcu_content.index("USB Type-C Controller(s) Firmware Version:\n")+2].split()[-1]
                numdot=re.compile("[0-9]+")
                if not numdot.match(firstPD) :
                    outputFile[0].at[i, "Reference Info"] = "N\A"
                #check secondPD is exit or not
                elif numdot.match(firstPD) and not numdot.match(secondPD) :
                    outputFile[0].at[i, "Reference Info"] = firstPD
                    continue
                else :
                    outputFile[0].at[i, "Reference Info"] = firstPD + "\n" + secondPD
                    continue
            except :
                pass
        elif i == "PCR[00] TPM 2.0 SHA256" or i == "PCR 0" :
            try :
                if SHA256_content :
                    try : # if SHA256UTF16le :
                        indexOfSHA = int(SHA256_content.index("TPM2_Startup: Return Code: 0x100\n"))+1
                        sha256 = SHA256_content[indexOfSHA:indexOfSHA+2]
                        sha256[0] = sha256[0][8:-3]
                        sha256[1] = sha256[1][8:-2]
                        sha256 = ''.join(sha256[0].split()) + ''.join(sha256[1].split())
                        outputFile[0].at[i, "Reference Info"] = sha256
                    except Exception as e : #else :
                        print(e)
                        print("Try to get pcr made of winpvt")
                        sha256 = re.compile(".*PCR Index 00:.*")
                        sha256 = list( filter( sha256.match, SHA256_content ) )[0].split()[-1]
                        outputFile[0].at[i, "Reference Info"] = sha256
                    continue
                else :
                    print("SHA256_content is empty!")
            except Exception as e :
                print("Get PCR ERROR MSG : ", e)
        elif (i == "FUR" or i == "HPBIOSUPDREC") and not isAMIPlatform:
            furV = ""
            if furP :
                furV = information_parser.GetFileVersion(furP)
            outputFile[0].at[i, "Reference Info"] = furV
            continue
        elif i == "SVN ver. Core" :
            if el_name :
                if ("R" in platform) or ("S" in platform) :     #G5/G6
                    svn_core = re.compile(".*HpCore\n")
                else : # ("" in platform)                       #G8 git
                    svn_core = re.compile(".*HpCorePvtBins\n")
                svn_core = list( filter( svn_core.match, el_content ) )[0].split()[1]
                outputFile[0].at[i, "Reference Info"] = svn_core
                continue
        elif i == "SVN ver. Chipset & PE" :
            if el_name :
                if isAMDPlatform :                              #AMD
                    svn_pi = re.compile(".*AMD\n")
                else :                                          #Intel
                    svn_pi = re.compile(".*HpIntelChipsetPkg\n")
                svn_pi = list( filter( svn_pi.match, el_content ) )[0].split()[1]
                outputFile[0].at[i, "Reference Info"] = svn_pi
                continue
        ###some should not get 
        elif i in {"Configurations", "1st", "2nd", "3rd", "Average"\
                    , "BIOS Initialization Duration(MS)", "Total Boot Duration(MS)"\
                    , "BiosInitTime(MS)", "DriverWakeTime(MS)"\
                    , "NOTE FOR THIS BIOS RELEASE", "TOOL REVISION"\
                    , "System POST TIME", "BIOS MODULE INFORMATION"\
                    , "BOOT TIME (ADK)", "S3 RESUME TIME"\
                    , "Known SI issues ready for retest with this release"\
                    , "Issue lists", "EC/SIO Functional changes"\
                    , "Configuration Table Information" \
                    , "Peripheral Information", "Processor ID" \
                    , "Intel IHV Information", "AMD IHV Information" \
                    , "FW Capsule Information" \
                    } :
            continue
        ##########AMD start
        elif i == "AMD Agesa PI" or i == "AMD Agesa code" :
            if amdz_name :
                agesaPI = re.compile("AGESA:.*")
                agesaPI = list( filter( agesaPI.match, amdz_content ) )[0].split()[-1] #agesaPI = agesaPI[0].split()[-1]
                outputFile[0].at[i, "Reference Info"] = agesaPI
                continue
        elif i == "PSP FW" :
            if amdz_name :
                PSPandSMU = re.compile("SMU:.*")
                PSPandSMU = list( filter( PSPandSMU.match, amdz_content ) )
                PSPandSMU = PSPandSMU[0].split()
                pspfw = PSPandSMU[2]
                if "(" in pspfw  :
                    pspfw = pspfw.split("(")[1][:-1]
                else :
                    pspfw = pspfw[2:]
                realPSPFW = ""
                for j in range(0,len(pspfw),2) :
                    if pspfw[j] == "0" and pspfw[j+1] == "0" :
                        realPSPFW = realPSPFW + "0."
                    elif pspfw[j] == "0" and pspfw[j+1] != "0" :
                        realPSPFW = realPSPFW + pspfw[j+1] + "."
                    else : 
                        realPSPFW = realPSPFW + pspfw[j] + pspfw[j+1] + "."
                realPSPFW = realPSPFW[:-1]
                outputFile[0].at[i, "Reference Info"] = realPSPFW
                continue
        elif i == "SMU FW" : #maybe dec to hex
            if amdz_name :
                smufw = PSPandSMU[1].split("(")[0]
                outputFile[0].at[i, "Reference Info"] = smufw
                continue
        elif i == "AMD Legacy VBIOS" or i == "AMD VBIOS" :
            if amdz_name :
                vBIOS = re.compile("VBIOS Info.*")
                vBIOS = list( filter( vBIOS.match, amdz_content ) )[0].split()[3][0:-1]
                outputFile[0].at[i, "Reference Info"] = vBIOS
                continue
        elif i == "AMD GOP EFI Driver" or i == "AMD GOP" :
            try :
                gOP = re.compile("Rev.*")
                gOP = list( filter( gOP.match, bcu_content[bcu_content.index("Video BIOS Version\n")+1].split() ) )
                gOP = gOP[0][4:-4]
                outputFile[0].at[i, "Reference Info"] = gOP
                continue
            except Exception:
                print("AMD GOP ERROR MSG : ", Exception)
        ##########AMD end
        ##########Intel start
        elif (i == "ME Firmware" or i == "Intel (R) Converged Security and Management Engine")\
              and not isAMIPlatform :
            try :
                mef = bcu_content[bcu_content.index("ME Firmware Version\n")+1].strip()
                mef = "Corporate  v"+mef
                outputFile[0].at[i, "Reference Info"] = mef
                continue
            except :
                pass
        elif i == "Reference Code" and not isAMIPlatform :
            try :
                rc = bcu_content[bcu_content.index("Reference Code Revision\n")+1].strip()
                outputFile[0].at[i, "Reference Info"] = rc
                continue
            except :
                pass
        elif i == "Intel GOP EFI Driver" and not isAMIPlatform :
            try :
                igop = bcu_content[bcu_content.index("Video BIOS Version\n")+1].split("[")[-1].split("]")[0]
                outputFile[0].at[i, "Reference Info"] = igop
                continue
            except :
                pass
        elif i == "GbE Version" and not isAMIPlatform and not isAMDPlatform :
            if binFile :
                with open(binFile, 'rb') as f:
                    f.seek(4106,0)
                    content1 = f.read(1)
                    content2 = f.read(1)
                if str(hex(ord(content2.decode()))).split("x")[-1] == "0" :
                    gbev = "00."+str(hex(ord(content1.decode()))).split("x")[-1]
                else :
                    gbev = str(hex(ord(content2.decode()))).split("x")[-1] \
                        +"."+str(hex(ord(content1.decode()))).split("x")[-1]
                    gbev = gbev[:-1]
                    
                outputFile[0].at[i, "Reference Info"] = gbev
            continue
        elif i == "AMD IHV Information" :
            isAMDBlock = True
        elif not i or i == "" or i == "\n" :
            isAMDBlock = False
        elif i == "Processor Microcode Patches" and not isAMDBlock and not isAMIPlatform :
            try :
                pm = bcu_content[bcu_content.index("Processor 1 MicroCode Revision\n")+1].strip()
                outputFile[0].at[i, "Reference Info"] = "0x"+pm
                if gCount["Microcode Version"] != 1 :
                    ishc = " \ " + gCount["Microcode Version"]
                    outputFile[0].at[i, "Reference Info"] += str(ishc)
                continue
            except :
                pass
        elif i == "MRC" :
            #print("IN MRC")
            #print(pdfTables[2])
            try :
                ################# here from pdf
                # tables[2]
                # "Unnamed: 0" -> Firmware Ingredient
                # "Unnamed: 1" -> Version 
                # "Unnamed: 2" -> Source (GIT/RDC) 
                # "Unnamed: 3" -> Changes
                #for i, df in enumerate(pdfTables, start=0):
                #    if "Firmware Ingredient" in df.columns:
                #        tNum = i
                #mrc = list(pdfTables[tNum].loc[pdfTables[tNum]["Unnamed: 0"]=="Silicon Initialization Code","Unnamed: 1"])
                #mrc = mrc[0].split()[-1]
                #mrc = mrc.split(")")[0]
                #print("mrc:",mrc)
                #outputFile[0].at[i, "Reference Info"] = str(mrc)
                #print(outputFile[0].at[i, "Reference Info"])
                ################# here from System Scope
                if gCount["Reference Code - MRC"] != 1 :
                    mrc = gCount["Reference Code - MRC"].split("(")[-1][:-1]
                    print("mrc : ",mrc)
                    outputFile[0].at[i, "Reference Info"] = str(mrc)
                continue
            except Exception as e :
                print(e)
                pass
        elif i == "ISH FW version" :
            try :
                if gCount["ISHC FW Version"] != 1 :
                    ishc = gCount["ISHC FW Version"].split("(")[-1][:-1]
                    outputFile[0].at[i, "Reference Info"] = str(ishc)
                continue
            except Exception as e :
                print(e)
                pass
        elif i == "ACM Module" :
            #print("IN ACM")
            #print(pdfTables[2])
            try :
                ################# here from pdf
                # tables[2]
                # "Unnamed: 0" -> Firmware Ingredient
                # "Unnamed: 1" -> Version 
                # "Unnamed: 2" -> Source (GIT/RDC) 
                # "Unnamed: 3" -> Changes
                #for i, df in enumerate(pdfTables, start=0):
                #    if "Firmware Ingredient" in df.columns:
                #        tNum = i
                #acmNum = pdfTables[tNum][pdfTables[tNum]["Unnamed: 0"]=="Intel® TXT and Intel® Boot Guard ACM* and SINIT"]
                #acmNum = acmNum.index[0]
                ## Intel® TXT and Intel® Boot Guard ACM* and SINIT in line 21, but acm in 20
                #acm = pdfTables[tNum].iloc[acmNum-1]["Unnamed: 1"]
                #print("acm:",acm)
                #outputFile[0].at[i, "Reference Info"] = str(acm)
                #print(outputFile[0].at[i, "Reference Info"])
                ################# here from System Scope
                if gCount["TXT ACM version"] != 1 :
                    acm = gCount["TXT ACM version"].split("(")[-1][:-1]
                    outputFile[0].at[i, "Reference Info"] = str(acm)
                continue
            except Exception as e :
                print(e)
                pass
        elif i == "SINIT ACM" :
            #rint("IN SINIT")
            #print(pdfTables[2])
            try :
                # tables[2]
                # "Unnamed: 0" -> Firmware Ingredient
                # "Unnamed: 1" -> Version 
                # "Unnamed: 2" -> Source (GIT/RDC) 
                # "Unnamed: 3" -> Changes
                for i, df in enumerate(pdfTables, start=0):
                    if "Firmware Ingredient" in df.columns:
                        tNum = i
                sinitN = pdfTables[tNum][pdfTables[tNum]["Unnamed: 0"]=="Intel® TXT and Intel® Boot Guard ACM* and SINIT"]
                sinitN = sinitN.index[0]
                # Intel® TXT and Intel® Boot Guard ACM* and SINIT in line 21, but sinit in 22
                sinit = pdfTables[tNum].iloc[sinitN-1]["Unnamed: 1"]
                print("sinit : ",sinit)
                outputFile[0].at[i, "Reference Info"] = str(sinit)
                #print(outputFile[0].at[i, "Reference Info"])
                continue
            except Exception as e :
                print(e)
                pass
        elif i == "PPAM" :
            #print("IN PPAM")
            #print(pdfTables[3])
            try :
                for i, df in enumerate(pdfTables, start=0):
                    if "Ingredient" in df.columns:
                        tNum = i
                ppam = list(pdfTables[tNum].loc[pdfTables[tNum]["Ingredients"]=="Intel® Platform Properties Assessment Module (PPAM)","Version"])
                ppam = ppam[0]
                print("ppam : ",ppam)
                outputFile[0].at[i, "Reference Info"] = str(ppam)
                #print(outputFile[0].at[i, "Reference Info"])
                continue
            except Exception as e :
                print(e)
                pass
        elif i == "ChipSetinit" :
            try :
                if me_name :
                    chipSetinit = re.compile(".*OEM Chipset Init Version.*")
                    chipSetinit = list( filter( chipSetinit.match, me_content ) )[0].split()[-1]
                    outputFile[0].at[i, "Reference Info"] = chipSetinit
                    continue
                elif gCount["OEM Chipset Init Version"] != 1 :
                    chipSetinit = gCount["OEM Chipset Init Version"]
                    outputFile[0].at[i, "Reference Info"] = str(chipSetinit)
            except :
                pass
        elif i == "NPHY FW version" or i == "NPHY FW  version" :
            try :
                if me_name :
                    nphy = re.compile(".*NPHY FW Version.*")
                    nphy = list( filter( nphy.match, me_content ) )[0].split()[-1]
                    outputFile[0].at[i, "Reference Info"] = nphy
                    continue
            except :
                pass
        elif i == "PMC" :
            try :
                if me_name :
                    pmc = re.compile(".*PMC FW Version.*")
                    pmc = list( filter( pmc.match, me_content ) )[0].split()[-1]
                    outputFile[0].at[i, "Reference Info"] = pmc
                    continue
                elif gCount["PMC FW Version"] != 1 :
                    pmc = gCount["PMC FW Version"].split("(")[-1][:-1]
                    outputFile[0].at[i, "Reference Info"] = str(pmc)
            except :
                pass
        elif i == "I225 Undi Driver" :
            try :
                if dri_name :
                    undi = re.compile(".*2.5G.*")
                    undi = list( filter( undi.match, dri_content ) )[0].split()[1]
                    undi = [undi[i:i+2] for i in range(0, len(undi), 2)]
                    undi = ".".join(str(int(i, 16)) for i in undi)
                    outputFile[0].at[i, "Reference Info"] = undi
                    continue
            except Exception as e :
                print(e)
                pass
        elif i == "PXE UEFI Driver" :
            try :
                if dri_name :
                    pxeUefi = re.compile(".*I219.*")
                    pxeUefi = list( filter( pxeUefi.match, dri_content ) )[0].split()[1]
                    pxeUefi = [pxeUefi[i:i+2] for i in range(0, len(pxeUefi), 2)]
                    pxeUefi = ".".join(str(int(i, 16)) for i in pxeUefi)
                    outputFile[0].at[i, "Reference Info"] = pxeUefi
                    continue
            except Exception as e :
                print(e)
                pass
        ######Intel U21/23 new release have some new items which need to try to get from Function Changes 
        ###### new U21/23 start
        elif i in { "Sprint", "Camera FW", "Touch controller FW", "Clickpad FW", "Fingerprint FW" \
                  , "RGB keyboard controller firmware version", "Boot Guard ACM" \
                  } and not isAMIPlatform :
            try :
                changes_content=list(outputFile[0].at["BIOS Functional changes", "Release Note Info"].split("\n"))
                #re.IGNORECASE that case unsensetive
                getFromchange = re.compile(".*"+i.split()[0]+".*", re.IGNORECASE) 
                getFromchange = list( filter( getFromchange.match, changes_content ) )
                if not getFromchange :
                    outputFile[0].at[i, "Reference Info"] = "Can not get from Function Changes!"
                elif len(getFromchange) == 1 :
                    outputFile[0].at[i, "Reference Info"] = getFromchange[0]
                else : # multy
                    outputFile[0].at[i, "Reference Info"] = "\n".join(getFromchange)
                #print(i+":", end="")
                #print(getFromchange)
                #print(outputFile[0].at[i, "Reference Info"])
                continue
                #print(changes_content)
            except Exception as e :
                print(e)
                pass
        ###### new U21/23 end
        ##########Intel end
        ##########AMI start
        elif isAMIPlatform :
            try :
                #print("AMI!! "+i)
                if i == "System BIOS Version" :
                    biosversion = bcu_content[bcu_content.index("BIOS Revision\n")+1].strip()[1:]
                    outputFile[0].at[i, "Reference Info"] = biosversion
                    continue
                elif i == "Build Date" :
                    try :
                        bdate = bcu_content[bcu_content.index("BIOS Date (ReadOnly)\n")+1].strip()
                    except :
                        try :
                            bdate = bcu_content[bcu_content.index("BIOS Date \n")+1].strip()
                        except :
                            bdate = bcu_content[bcu_content.index("BIOS Date\n")+1].strip()
                    outputFile[0].at[i, "Reference Info"] = bdate
                    continue
                elif i == "CHECKSUM" :
                    if binFile :
                        with open(binFile, 'rb') as f:
                            content = f.read()
                            binary_sum = sum(bytearray(content))
                            binary_sum = hex(binary_sum & 0xFFFFFFFF)
                            f.close()
                        binary_sum = binary_sum.split("x")[-1].upper()
                        outputFile[0].at[i, "Reference Info"] = binary_sum
                    continue
                elif i == "FUR" :
                    if ami_furP :
                        furV = information_parser.GetFileVersion(ami_furP)
                    outputFile[0].at[i, "Reference Info"] = furV
                    continue
                elif i == "ME Firmware" :
                    try :
                        me_name = re.compile(r"ME_[0-9\.]+\.bin") 
                        os.chdir(release_dir+"\\METools\\FWUpdate\\MEFW")
                        me_dir = os.listdir( os.getcwd() ) #list
                        me_name = list( filter( me_name.match, me_dir ) )
                        os.chdir(new_dir)
                    except :
                        print("Get AMI ME folder fail !")
                    if len(me_name) > 0 :
                        mef = me_name[0][3:-4]+"_Consumer"  #ME_[0-9\.]+\.bin
                    else :
                        print("\nCan not find .\\METools\\FWUpdate\\MEFW\\ME_\{version\}.bin !")
                    outputFile[0].at[i, "Reference Info"] = mef
                    continue
                elif i == "GbE Version" :
                    continue
                else :
                    pass
            except :
                pass
        ##########AMI end
        ##########G4 start
        elif i == "System BIOS" and isAMDG4Platform :
            try:
                bversion = bcu_content[bcu_content.index("System BIOS Version\n")+1].split()[2]
                outputFile[0].at[i, "Reference Info"] = "Ver " + bversion
                continue
            except :
                pass
        ##########G4 end
        ##########new U21/U23 start
        elif ( i == "System BIOS" or i == "HP System Firmware" ) \
              and \
             ( not isAMDPlatform and not isAMIPlatform ) : 
            try:
                bversion = bcu_content[bcu_content.index("System BIOS Version\n")+1].split()[2]
                outputFile[0].at[i, "Reference Info"] = bversion
                continue
            except :
                pass
        ##########new U21/U23 end
        else : 
            pass
        print("Can not find : " + str(i) )
        outputFile[0].at[i, "Reference Info"] = "N/A"

outputFile[0]["Release Note Info"].fillna(value="N/A",inplace=True)
outputFile[0]["Reference Info"].fillna(value="N/A",inplace=True)
#=========Resault
#compare
for i in rRowInfoName:
    if type(i) != str or i == "TOOL REVISION" or i == "NOTE FOR THIS BIOS RELEASE"\
        or i == "System POST TIME" or i == "BOOT TIME (ADK)" \
        or i == "S3 RESUME TIME" or i == "BIOS MODULE INFORMATION" \
        or i == "EC/SIO Functional changes" or i == "Configuration Table Information" \
        or i == "Peripheral Information" or i == "Processor ID" \
        or i == "Intel IHV Information" or i == "AMD IHV Information" \
        or i == "FW Capsule Information" or i == "Known SI issues ready for retest with this release" \
        : # skip nan
        continue
    if str(outputFile[0].at[i, "Release Note Info"]) == "N/A" \
        and str(outputFile[0].at[i, "Reference Info"]) == "N/A" :
        outputFile[0].at[i, "Result"] = "Ignore"  #Both N/A
    elif i == "Build Date" :
        if set(str(outputFile[0].at[i, "Release Note Info"]).split("/")) \
        == set(str(outputFile[0].at[i, "Reference Info"]).split("/")) :
            outputFile[0].at[i, "Result"] = "V" 
        else :
            outputFile[0].at[i, "Result"] = "X" 
    elif i == "Processor Microcode Patches" :
        if str(outputFile[0].at[i, "Reference Info"]).split("x") \
        in str(outputFile[0].at[i, "Release Note Info"]).split("x") : # 0x123 or 0x0123 "in" 0x0123 is both OK
            outputFile[0].at[i, "Result"] = "V" 
        else :
            outputFile[0].at[i, "Result"] = "X" 
    elif i == "PCR[00] TPM 2.0 SHA256" or i == "PCR 0" :
        if ''.join(str(outputFile[0].at[i, "Release Note Info"]).split()) \
        == ''.join(str(outputFile[0].at[i, "Reference Info"]).split()) :
            outputFile[0].at[i, "Result"] = "V" 
        else :
            outputFile[0].at[i, "Result"] = "X" 
    elif i in { "Sprint", "Camera FW", "Touch controller FW", "Clickpad FW", "Fingerprint FW" \
            , "RGB keyboard controller firmware version", "Boot Guard ACM" } :
        if str(outputFile[0].at[i, "Release Note Info"]) in str(outputFile[0].at[i, "Reference Info"]) :
            outputFile[0].at[i, "Result"] = "V"  
        else :
            outputFile[0].at[i, "Result"] = "X"
    elif str(outputFile[0].at[i, "Release Note Info"]) == str(outputFile[0].at[i, "Reference Info"]) :
        outputFile[0].at[i, "Result"] = "V"  
    else :
        outputFile[0].at[i, "Result"] = "X"
#clean na
for i in range(len(rRowInfoName) ):
    if outputFile[0].iat[i,2] != outputFile[0].iat[i,2] :
        outputFile[0].iat[i,1] = " "
        outputFile[0].iat[i,0] = " "
#protect excel
# current_date = datetime.now().strftime("%Y-%m-%d")
# current_date_bytes = current_date.encode("utf-8")
# hashed_key = hashlib.sha256(current_date_bytes).hexdigest()
# print(hashed_key)
# print(outputFile[0])
# print(type(outputFile[0]))
# outputFile[0].loc[1, "Release Note Info"].protection = openpyxl.styles.Protection(locked=True)
# for row in rRowInfoName:
#     print(row)
#     for col in ["Release Note Info", "Reference Info", "Reference Info"]:
#         print(col, end=" ")
#         cell_address = f'{col}{row}'
#         cell = outputFile_PlatformHistory[cell_address]

#         cell.protection = openpyxl.styles.Protection(locked=True)
# outputFile[0].security.set_workbook_password(hashed_key)
#====Safe file
outputFile[0] = outputFile[0]
outputFile_PlatformHistory.to_excel(writer,sheet_name="PlatformHistory")
writer.close()
print("\nComparison completed successfully!\n")
#====Gray line
print("Start gray line~")
os.chdir(new_dir)
workbook = openpyxl.load_workbook(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
sheet = workbook["PlatformHistory"]
grayList = []
grayUnder = "System POST TIME"
underFlag = False
for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row):
    tempG = row[0].value
    if tempG == grayUnder:
        underFlag = True
    if tempG in grayList or underFlag or row[3].value == "Ignore" :
        fill = openpyxl.styles.PatternFill(start_color='808080', end_color='808080', fill_type='solid')
        for cell in row:
            cell.fill = fill
workbook.save(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
"""
row_index = None
for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row):
    for cell in row:
        if cell.value == goal :
            row_index = cell.row
            break
if row_index is not None:
    # gray
    fill = openpyxl.styles.PatternFill(start_color="808080", end_color="808080", fill_type="solid")
    for row in sheet.iter_rows(min_row=row_index):
        for cell in row:
            cell.fill = fill
    workbook.save(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
else:
    print("Can not find the \"" + goal + "\" to gray the line under it!")
"""    
print("End gray line~")
#====Set width
print("Star set width and lock form~")
workbook = openpyxl.load_workbook(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
sheet = workbook['PlatformHistory']
sheet.column_dimensions['A'].width = 55
sheet.column_dimensions['B'].width = 28
sheet.column_dimensions['C'].width = 28
#====Get date_sha to lock excel
current_date = datetime.now().strftime("%Y-%m-%d")
current_date_bytes = current_date.encode("utf-8")
hashed_key = hashlib.sha256(current_date_bytes).hexdigest()
#print(hashed_key)
sheet.protection.sheet = True
sheet.protection.password = str(hashed_key)[0:3]
sheet.protection.enable()
print("End set width and lock form~")
workbook.save(str(goal_platform)+"_"+str(goal_version)+"_result_RN.xlsx")
os.system("pause")